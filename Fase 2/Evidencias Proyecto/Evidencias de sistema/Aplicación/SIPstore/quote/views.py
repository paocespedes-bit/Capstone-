from django.shortcuts import render
from sipstore import settings
from store.models import PanelSIP, Categoria
from django.db.models import Q
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse
import json
from django.core.mail import EmailMessage
from django.template.loader import render_to_string
from io import BytesIO
import pandas as pd
from docx import Document
from cart.carrito import Carrito
from store.models import PanelSIP
from django.contrib.contenttypes.models import ContentType
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def quote(request):
    
    print(settings.SENDGRID_API_KEY)
    modulos = {
        # !Ojo con las categorias tienen que tener estos mismos nombres
        'piso': 'Piso',
        'cielo': 'Cielo',
        'muros_int': 'Muros Internos',
        'muros_ext': 'Muros Externos', 
        
    }

    contexto = {}

    for key, nombre_categoria in modulos.items():
        categoria = Categoria.objects.filter(nombre__iexact=nombre_categoria).first()

        if categoria:
            paneles = (
                PanelSIP.objects.filter(categorias__id=categoria.id)
                .filter(
                    Q(inventario__modo_stock='stock', inventario__disponible__gt=0)
                    | Q(inventario__modo_stock='pedido')
                )
                .distinct()[:5]
            )
            #!Despues borrar
            print(f"==============================")
            print(f"Categoría buscada: {nombre_categoria}")
            print(f"✅ Categoría encontrada: {categoria.nombre} (id={categoria.id})")
            print(f"➡️ Paneles con esa categoría: {PanelSIP.objects.filter(categorias__id=categoria.id).count()}")
            print(f"➡️ Paneles con inventario válido: {paneles.count()}")
            print(f"==============================")
            #!======================================================
            contexto[key] = paneles
        else:
            print(f"Categoría buscada: {nombre_categoria} ❌ No se encontró la categoría.") #!Despues borrar
            contexto[key] = []
            
    request.session.flush()
    return render(request, 'quote.html', contexto)


@csrf_exempt
def save_form(request):
    if request.method == "POST":
        data = json.loads(request.body.decode("utf-8"))
        module = data.get("module")
        form_data = data.get("data")

        if not module or not form_data:
            return JsonResponse({"error": "Datos incompletos"}, status=400)

        session_forms = request.session.get("quote_forms", {})

        if module not in session_forms:
            session_forms[module] = []

        # --- Evitar duplicados y sobrescribir si ya existe ---
        updated = False
        for i, existing_form in enumerate(session_forms[module]):
            if existing_form == form_data:
                # Mismo formulario, no hacer nada
                return JsonResponse({"message": "Formulario ya guardado sin cambios"})
            # Si se parece (mismo tipoPanel y dimensiones similares), sobrescribir
            if (
                existing_form.get("tipoPanel") == form_data.get("tipoPanel")
                and existing_form.get("largo") == form_data.get("largo")
                and existing_form.get("ancho") == form_data.get("ancho")
            ):
                session_forms[module][i] = form_data
                updated = True
                break

        if not updated:
            session_forms[module].append(form_data)

        request.session["quote_forms"] = session_forms
        request.session.modified = True

        msg = "Formulario actualizado" if updated else "Formulario guardado temporalmente"
        return JsonResponse({"message": msg})

@csrf_exempt
def calcular_materiales(request):
    if request.method == "POST":
        quote_data = request.session.get("quote_forms", {})
        if not quote_data:
            return JsonResponse({"html": "<p class='text-center text-muted'>No hay datos para calcular</p>"})

        resultados = []
        total_general = 0

        muros = {}
        aberturas = []

        for modulo, forms in quote_data.items():
            for form in forms:
                try:
                    if modulo in ["muros-int", "muros-ext"]:
                        largo = float(form.get("largo") or 0)
                        alto = float(form.get("alto") or 0)
                        muro_id = form.get("id") or f"{modulo}-{len(muros.get(modulo, {})) + 1}"
                        area = largo * alto
                        muros.setdefault(modulo, {})[muro_id] = {
                            "largo": largo,
                            "alto": alto,
                            "area": area,
                        }

                    elif modulo == "abertura":
                        largo = float(form.get("largo") or 0)
                        ancho = float(form.get("ancho") or 0)
                        tipo_muro = form.get("tipoMuro", "").strip().lower()
                        muro_id = form.get("muroId")
                        area_abertura = largo * ancho

                        if tipo_muro in ["muro interior", "interior"]:
                            tipo_modulo = "muros-int"
                        else:
                            tipo_modulo = "muros-ext"

                        aberturas.append({
                            "tipo_modulo": tipo_modulo,
                            "muro_id": muro_id,
                            "area": area_abertura,
                        })

                except Exception as e:
                    print(f"Error registrando formulario {modulo}: {e}")

        for abertura in aberturas:
            tipo_modulo = abertura["tipo_modulo"]
            muro_id = abertura["muro_id"]
            area = abertura["area"]

            if tipo_modulo in muros and muro_id in muros[tipo_modulo]:
                muros[tipo_modulo][muro_id]["area"] = max(
                    0, muros[tipo_modulo][muro_id]["area"] - area
                )

        for modulo, forms in quote_data.items():
            for form in forms:
                try:
                    largo = float(form.get("largo") or 0)
                    ancho = float(form.get("ancho") or 0)
                    panel_id = int(form.get("tipoPanel") or 0)

                    if not (largo > 0 and ancho > 0 and panel_id):
                        continue

                    panel = PanelSIP.objects.get(id=panel_id)
                    imagen_rel = panel.imagenes.first()
                    imagen_url = imagen_rel.imagen.url if imagen_rel else "/static/img/SinImagen.png"

                    # --- Cálculo actualizado (medidas en metros) ---
                    area_total = largo * ancho

                    if modulo in ["muros-int", "muros-ext"]:
                        for muro_id, muro_data in muros[modulo].items():
                            if muro_data["largo"] == largo and muro_data["alto"] == ancho:
                                area_total = muro_data["area"]
                                break

                    # Ahora el área del panel está directamente en m²
                    area_panel = float(panel.largo) * float(panel.ancho)

                    if area_panel <= 0:
                        area_panel = 2.88  # valor por defecto si el panel no tiene datos válidos

                    cantidad = max(1, round(area_total / area_panel))
                    total = cantidad * float(panel.precio_actual)
                    total_general += total

                    resultados.append({
                        "nombre": panel.nombre,
                        "imagen": imagen_url,
                        "solicitado_por": modulo,
                        "area": round(area_total, 2),
                        "cantidad": cantidad,
                        "valor": round(panel.precio_actual, 2),
                        "total": round(total, 2),
                    })

                except PanelSIP.DoesNotExist:
                    print(f"Panel con ID {form.get('tipoPanel')} no existe.")
                except Exception as e:
                    print(f"Error calculando {modulo}: {e}")

        html = render(request, "partials/calculo_resultado.html", {
            "resultados": resultados,
            "total_general": total_general
        }).content.decode("utf-8")

        request.session["productos_calculo"] = resultados
        request.session.modified = True

        return JsonResponse({"html": html})

@csrf_exempt
def limpiar_calculo(request):
    request.session["quote_forms"] = {}
    request.session.modified = True
    return JsonResponse({"message": "Cálculo limpiado"})



@csrf_exempt
def descargar_cotizacion(request, formato):
    """
    Permite descargar la cotización actual en .txt, .docx o .xlsx
    """
    quote_data = request.session.get("quote_forms", {})
    if not quote_data:
        return JsonResponse({"error": "No hay datos para descargar"}, status=400)

    resultados = []
    total_general = 0

    for modulo, forms in quote_data.items():
        for form in forms:
            try:
                largo = float(form.get("largo") or 0)
                ancho = float(form.get("ancho") or 0)
                panel_id = int(form.get("tipoPanel") or 0)

                if not (largo > 0 and ancho > 0 and panel_id):
                    continue

                panel = PanelSIP.objects.get(id=panel_id)

                area_total = largo * ancho
                area_panel = float(panel.largo) * float(panel.ancho)
                if area_panel <= 0:
                    area_panel = 2.88

                cantidad = max(1, round(area_total / area_panel))
                total = cantidad * float(panel.precio_actual)
                total_general += total

                resultados.append({
                    "Panel": panel.nombre,
                    "Módulo": modulo,
                    "Área (m²)": round(area_total, 2),
                    "Cantidad": cantidad,
                    "Valor Unitario": round(panel.precio_actual, 2),
                    "Total": round(total, 2),
                })

            except PanelSIP.DoesNotExist:
                print(f"⚠️ Panel con ID {form.get('tipoPanel')} no existe.")
                continue
            except Exception as e:
                print(f"⚠️ Error en módulo {modulo}: {e}")
                continue

    if formato == "txt":
        buffer = BytesIO()
        contenido = "COTIZACIÓN PANEL SIP\n\n"
        for r in resultados:
            contenido += f"{r['Panel']} ({r['Módulo']}): {r['Cantidad']} panel(es) - Total: ${r['Total']}\n"
        contenido += f"\nTOTAL GENERAL: ${round(total_general, 2)}"
        buffer.write(contenido.encode("utf-8"))
        buffer.seek(0)

        response = HttpResponse(buffer, content_type="text/plain")
        response["Content-Disposition"] = 'attachment; filename="cotizacion.txt"'
        return response

    elif formato == "doc":
        doc = Document()
        doc.add_heading("Cotización Panel SIP", 0)
        for r in resultados:
            doc.add_paragraph(
                f"{r['Panel']} ({r['Módulo']}): {r['Cantidad']} panel(es) - Total: ${r['Total']}"
            )
        doc.add_paragraph(f"\nTOTAL GENERAL: ${round(total_general, 2)}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = 'attachment; filename="cotizacion.docx"'
        return response

    elif formato == "xlsx":
        df = pd.DataFrame(resultados)
        df.loc[len(df.index)] = ["", "", "", "", "TOTAL GENERAL", round(total_general, 2)]

        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name="Cotizacion")
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = 'attachment; filename="cotizacion.xlsx"'
        return response

    else:
        return JsonResponse({"error": "Formato no soportado"}, status=400)

    
    
@csrf_exempt
def agregar_al_carrito(request):
    if request.method == 'POST':
        productos = request.session.get("productos_calculo", [])

        if not productos:
            return JsonResponse({"success": False, "message": "No hay productos para agregar al carrito."})

        carrito = Carrito(request)

        
        content_type = ContentType.objects.get_for_model(PanelSIP)

        for item in productos:
            try:
                producto = PanelSIP.objects.filter(nombre=item["nombre"]).first()
                if producto:
                    carrito.agregar({
                        "id": producto.id,
                        "nombre": producto.nombre,
                        "precio_actual": producto.precio_actual,
                        "content_type_id": content_type.id,
                    }, item["cantidad"])
            except Exception as e:
                print(f"⚠️ Error al agregar {item['nombre']} al carrito: {e}")
                continue

        
        del request.session["productos_calculo"]

        return JsonResponse({"success": True, "message": "Productos agregados al carrito correctamente."})

    return JsonResponse({"success": False, "message": "Método no permitido."})

def generar_pdf_cotizacion(resultados, total_general):
    """
    Genera el PDF de la cotización usando ReportLab. 
    Devuelve los bytes del PDF para adjuntarlos a un correo.
    
    IMPORTANTE: Esta función incluye manejo de errores para evitar el 
    Error 500 si hay problemas con los datos de entrada.
    """
    try:
        buffer = BytesIO()
        # Inicializa el lienzo (canvas) con el buffer y el tamaño de página
        p = canvas.Canvas(buffer, pagesize=letter)

        p.setFont("Helvetica-Bold", 14)
        p.drawString(50, 750, "COTIZACIÓN PANEL SIP")

        p.setFont("Helvetica", 10)
        y = 720

        # --- Iteración y Dibujo de Resultados ---
        for r in resultados:
            # Línea de detalle principal
            p.drawString(50, y, f"{r['nombre']} ({r['solicitado_por']})")
            y -= 15
            
            # Línea de detalle secundario (asegúrate que los valores sean strings o números)
            p.drawString(60, y, f"Área: {r['area']} m² | Cantidad: {r['cantidad']} | Valor unit: ${r['valor']} | Total: ${r['total']}")
            y -= 25

            # Salto de página si el contenido está muy abajo
            if y < 50:
                p.showPage()
                y = 750
        
        # --- Total General ---
        p.setFont("Helvetica-Bold", 12)
        # Asegúrate de que total_general sea un número
        p.drawString(50, y, f"TOTAL GENERAL: ${round(total_general, 2)}")

        # Finaliza el documento
        p.showPage()
        p.save()
        
        # Obtiene los bytes del buffer
        pdf = buffer.getvalue()
        buffer.close()
        
        # ¡IMPORTANTE! El nombre de la variable de retorno es 'pdf'
        return pdf
    
    except KeyError as e:
        # Maneja el error si falta alguna clave en el diccionario 'r'
        # Esto sucede si, por ejemplo, falta 'nombre' en un resultado.
        print(f"Error 500: Clave faltante en 'resultados': {e}")
        # Relanza la excepción para que el framework la capture y muestre el stack trace
        raise e
    except Exception as e:
        # Maneja cualquier otro error (ej. ReportLab, tipo de datos inválido)
        print(f"Error 500: Fallo general en generar_pdf_cotizacion: {e}")
        raise e
    
def enviar_cotizacion_por_correo(correo, resultados, total_general):
    pdf_bytes = generar_pdf_cotizacion(resultados, total_general)

    email = EmailMessage(
        subject="Tu cotización de Paneles SIP",
        body="Adjuntamos la cotización solicitada.\n\nGracias por cotizar con nosotros.",
        from_email="no-reply@tuempresa.cl",
        to=[correo],
    )

    email.attach("cotizacion.pdf", pdf_bytes, "application/pdf")
    email.send()

@csrf_exempt
def enviar_cotizacion(request):
    if request.method != "POST":
        return JsonResponse({"error": "Método no permitido"}, status=405)

    # 1. Leer y validar JSON
    try:
        data = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"error": "JSON inválido"}, status=400)

    correo = data.get("correo")
    if not correo:
        return JsonResponse({"error": "Correo no recibido"}, status=400)

    # 2. Obtener datos de la sesión
    productos = request.session.get("productos_calculo")
    if not productos:
        return JsonResponse({"error": "No hay datos de cotización"}, status=400)

    # 3. Calcular total
    total_general = sum(item["total"] for item in productos)

    # 4. Generar PDF
    # La generación del PDF utiliza tu función previamente revisada con manejo de errores.
    # Si esta función falla (por KeyError, ReportLab, etc.), se relanzará una excepción aquí.
    try:
        # Llamamos a la función dedicada para generar los bytes del PDF
        pdf_bytes = generar_pdf_cotizacion(productos, total_general)
    except Exception as e:
        print(f"ERROR: Fallo al generar el PDF: {e}")
        # Devolvemos 500 si la generación del PDF falla (pero el error interno ya está loggeado)
        return JsonResponse({"error": "Fallo al generar el archivo PDF de la cotización."}, status=500)

    # 5. Enviar email con PDF adjunto (Nueva sección de manejo de errores)
    try:
        email = EmailMessage(
            subject="Tu cotización de Paneles SIP",
            body="Adjuntamos tu cotización en formato PDF.\n\nGracias por cotizar con nosotros.",
            from_email=settings.DEFAULT_FROM_EMAIL, # Es una buena práctica definir from_email
            to=[correo],
        )
        # Adjuntamos los bytes del PDF
        email.attach("cotizacion.pdf", pdf_bytes, "application/pdf")
        email.send()

    except Exception as e:
        # Captura errores de conexión SMTP, autenticación o configuración de correo
        print(f"ERROR: Fallo al enviar el correo (SMTP/Configuración): {e}")
        # Devolvemos 500 al cliente con un mensaje útil
        return JsonResponse({"error": "Fallo al conectar o enviar el correo. Revisa la configuración SMTP del servidor."}, status=500)

    # 6. Respuesta de éxito
    return JsonResponse({"message": "Cotización enviada correctamente al correo."})

